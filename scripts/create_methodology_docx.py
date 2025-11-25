#!/usr/bin/env python3
"""
Скрипт для создания методологического документа DOCX
Методика расчёта нормативной численности ИТР
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Установить цвет фона ячейки"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_methodology_document():
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # === ТИТУЛЬНАЯ СТРАНИЦА ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('\n\n\n\n\n')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('МЕТОДИКА')
    run.bold = True
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('расчёта нормативной численности\nинженерно-технических работников (ИТР)\nна строительных проектах')
    run.font.size = Pt(16)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('\n\n\nВерсия 2.0')
    run.font.size = Pt(14)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run('\n\n2025 год')
    run.font.size = Pt(12)

    doc.add_page_break()

    # === ОГЛАВЛЕНИЕ ===
    doc.add_heading('СОДЕРЖАНИЕ', level=1)

    toc_items = [
        ('1. Область применения', 3),
        ('2. Термины и определения', 3),
        ('3. Классификация должностей ИТР', 4),
        ('4. Определение масштаба проекта', 5),
        ('5. Методика расчёта обязательных должностей', 6),
        ('6. Методика расчёта условных должностей', 8),
        ('7. Коэффициенты по масштабу проекта', 10),
        ('8. Примеры расчёта', 11),
        ('Приложение А. Условные вопросы калькулятора', 13),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f'{item}')
        p.add_run('\t' * 8 + str(page))

    doc.add_page_break()

    # === 1. ОБЛАСТЬ ПРИМЕНЕНИЯ ===
    doc.add_heading('1. Область применения', level=1)

    doc.add_paragraph(
        'Настоящая методика устанавливает порядок расчёта нормативной численности '
        'инженерно-технических работников (ИТР) для строительных проектов компании.'
    )

    doc.add_paragraph(
        'Методика предназначена для:'
    )

    items = [
        'планирования потребности в ИТР на этапе подготовки проекта;',
        'оценки оптимальности текущей численности ИТР;',
        'обоснования изменения штатного расписания;',
        'сравнительного анализа эффективности управления проектами.'
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        '\nМетодика разработана на основе анализа данных 74 проектов за период '
        'январь-октябрь 2025 года.'
    )

    # === 2. ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ ===
    doc.add_heading('2. Термины и определения', level=1)

    terms = [
        ('ИТР', 'Инженерно-технические работники — сотрудники, осуществляющие '
                'организацию и управление производственным процессом на строительной площадке.'),
        ('Рабочие', 'Производственный персонал, непосредственно выполняющий '
                   'строительно-монтажные работы.'),
        ('Масштаб проекта', 'Категория проекта (S/M/L/XL), определяемая по '
                           'численности рабочих.'),
        ('Обязательная должность', 'Должность, которая должна быть укомплектована '
                                   'на любом проекте независимо от его специфики.'),
        ('Условная должность', 'Должность, необходимость которой определяется '
                               'наличием специфических факторов на проекте.'),
        ('Коэффициент K', 'Нормативное количество рабочих на одного специалиста '
                         'данной должности.'),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Термин'
    hdr_cells[1].text = 'Определение'
    set_cell_shading(hdr_cells[0], 'D9E2F3')
    set_cell_shading(hdr_cells[1], 'D9E2F3')

    for term, definition in terms:
        row_cells = table.add_row().cells
        row_cells[0].text = term
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = definition

    # === 3. КЛАССИФИКАЦИЯ ДОЛЖНОСТЕЙ ИТР ===
    doc.add_heading('3. Классификация должностей ИТР', level=1)

    doc.add_paragraph(
        'Все должности ИТР разделены на три категории:'
    )

    doc.add_heading('3.1. Обязательные должности (6 групп)', level=2)

    doc.add_paragraph(
        'Должности, которые должны быть на каждом проекте:'
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    headers = ['№', 'Группа должностей', 'Основание расчёта', 'Минимум']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'C6EFCE')

    mandatory = [
        ('1', 'Руководитель проекта', 'Всегда 1 на проект', '1'),
        ('2', 'Производитель работ', 'По численности рабочих', '1'),
        ('3', 'Мастер', 'По численности рабочих', '1'),
        ('4', 'Специалист по охране труда', 'По законодательству (1 на 50 чел)', '1'),
        ('5', 'Специалист по общим вопросам', 'По численности ИТР', '1'),
        ('6', 'Кладовщик / Специалист ОМТС', 'По численности рабочих', '1'),
    ]

    for row_data in mandatory:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.add_heading('3.2. Условные должности (5 групп)', level=2)

    doc.add_paragraph(
        'Должности, необходимость которых определяется спецификой проекта:'
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    headers = ['№', 'Группа должностей', 'Условие включения']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'FFEB9C')

    conditional = [
        ('7', 'Водитель / Машинист / Механик', 'Наличие автотранспорта на проекте'),
        ('8', 'Инспектор строительных лесов', 'Использование строительных лесов'),
        ('9', 'Специалист по сопровождению групп', 'Привлечение иностранных рабочих'),
        ('10', 'Сотрудник службы безопасности', 'Требования охраны объекта'),
        ('11', 'Инженер-конструктор', 'Проектные работы на площадке'),
    ]

    for row_data in conditional:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.add_heading('3.3. Исключаемые должности (1 группа)', level=2)

    doc.add_paragraph(
        'Должности, НЕ включаемые в расчёт нормативной численности ИТР:'
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    headers = ['№', 'Группа должностей', 'Причина исключения']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'FFC7CE')

    row_cells = table.add_row().cells
    row_cells[0].text = '12'
    row_cells[1].text = 'Инструктор / Преподаватель'
    row_cells[2].text = 'Обучение — отдельный процесс, не относится к производственному ИТР'

    doc.add_page_break()

    # === 4. ОПРЕДЕЛЕНИЕ МАСШТАБА ПРОЕКТА ===
    doc.add_heading('4. Определение масштаба проекта', level=1)

    doc.add_paragraph(
        'Масштаб проекта определяется по среднемесячной численности рабочих:'
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    headers = ['Масштаб', 'Код', 'Численность рабочих', 'Характеристика']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'D9E2F3')
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    scales = [
        ('Малый', 'S', 'до 50 чел', 'Небольшой объект, 1-2 участка работ'),
        ('Средний', 'M', '50–150 чел', 'Стандартный объект, несколько участков'),
        ('Крупный', 'L', '150–300 чел', 'Большой объект, множество участков'),
        ('Очень крупный', 'XL', '300+ чел', 'Мега-проект, сложная структура'),
    ]

    for row_data in scales:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            if i == 2:
                row_cells[i].paragraphs[0].runs[0].bold = True

    doc.add_paragraph(
        '\nПримечание: При пограничных значениях численности рекомендуется '
        'выбирать бо́льший масштаб проекта.'
    ).italic = True

    # === 5. МЕТОДИКА РАСЧЁТА ОБЯЗАТЕЛЬНЫХ ДОЛЖНОСТЕЙ ===
    doc.add_heading('5. Методика расчёта обязательных должностей', level=1)

    doc.add_heading('5.1. Руководитель проекта', level=2)
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('1 на проект (константа)')
    doc.add_paragraph('Руководитель проекта назначается на каждый проект независимо от его масштаба.')

    doc.add_heading('5.2. Производитель работ (прораб)', level=2)
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('ceil(Рабочие / K_прораб)')
    doc.add_paragraph('Где K_прораб — коэффициент, зависящий от масштаба проекта (см. раздел 7).')

    doc.add_heading('5.3. Мастер', level=2)
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('ceil(Рабочие / K_мастер)')
    doc.add_paragraph('Где K_мастер — коэффициент, зависящий от масштаба проекта (см. раздел 7).')

    doc.add_heading('5.4. Специалист по охране труда', level=2)
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('ceil(Рабочие / 50)')
    doc.add_paragraph(
        'Норматив установлен в соответствии с требованиями законодательства РФ '
        'об охране труда (1 специалист на 50 работников).'
    )

    doc.add_heading('5.5. Специалист по общим вопросам', level=2)
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('ceil(Всего_ИТР / 15)')
    doc.add_paragraph('Административная поддержка рассчитывается от общей численности ИТР.')

    doc.add_heading('5.6. Кладовщик / Специалист ОМТС', level=2)
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('ceil(Рабочие / K_склад)')
    doc.add_paragraph('Где K_склад — коэффициент, зависящий от масштаба проекта (см. раздел 7).')

    doc.add_page_break()

    # === 6. МЕТОДИКА РАСЧЁТА УСЛОВНЫХ ДОЛЖНОСТЕЙ ===
    doc.add_heading('6. Методика расчёта условных должностей', level=1)

    doc.add_paragraph(
        'Условные должности включаются в расчёт только при наличии соответствующих '
        'факторов на проекте.'
    )

    doc.add_heading('6.1. Водитель / Машинист / Механик', level=2)
    p = doc.add_paragraph()
    p.add_run('Условие: ').bold = True
    p.add_run('На проекте используется автотранспорт')
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('ceil(Количество_единиц_техники / 3)')

    doc.add_heading('6.2. Инспектор строительных лесов', level=2)
    p = doc.add_paragraph()
    p.add_run('Условие: ').bold = True
    p.add_run('На проекте используются строительные леса')
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('ceil(Площадь_лесов_м² / 1000)')

    doc.add_heading('6.3. Специалист по сопровождению групп', level=2)
    p = doc.add_paragraph()
    p.add_run('Условие: ').bold = True
    p.add_run('На проекте привлекаются иностранные рабочие')
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('ceil(Количество_иностранных_рабочих / 100)')

    doc.add_heading('6.4. Сотрудник службы безопасности', level=2)
    p = doc.add_paragraph()
    p.add_run('Условие: ').bold = True
    p.add_run('Требуется охрана объекта')
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('Количество_постов × (4 при круглосуточной охране, иначе 2)')

    doc.add_heading('6.5. Инженер-конструктор', level=2)
    p = doc.add_paragraph()
    p.add_run('Условие: ').bold = True
    p.add_run('На площадке выполняются проектные работы')
    p = doc.add_paragraph()
    p.add_run('Формула: ').bold = True
    p.add_run('ceil(Рабочие / 100)')

    doc.add_page_break()

    # === 7. КОЭФФИЦИЕНТЫ ПО МАСШТАБУ ПРОЕКТА ===
    doc.add_heading('7. Коэффициенты по масштабу проекта', level=1)

    doc.add_paragraph(
        'Коэффициент K определяет нормативное количество рабочих на одного '
        'специалиста соответствующей должности.'
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    headers = ['Масштаб', 'Рабочих', 'K_прораб', 'K_мастер', 'K_склад']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'D9E2F3')
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    coefficients = [
        ('Малый (S)', 'до 50', '15', '10', '30'),
        ('Средний (M)', '50–150', '20', '12', '40'),
        ('Крупный (L)', '150–300', '25', '15', '50'),
        ('Очень крупный (XL)', '300+', '30', '18', '60'),
    ]

    for row_data in coefficients:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.add_paragraph(
        '\nИнтерпретация коэффициентов:'
    )

    interpretations = [
        'K_прораб = 25 означает: 1 производитель работ на 25 рабочих',
        'K_мастер = 15 означает: 1 мастер на 15 рабочих',
        'K_склад = 50 означает: 1 кладовщик на 50 рабочих',
    ]
    for item in interpretations:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # === 8. ПРИМЕРЫ РАСЧЁТА ===
    doc.add_heading('8. Примеры расчёта', level=1)

    doc.add_heading('8.1. Пример: Крупный проект (200 рабочих)', level=2)

    doc.add_paragraph('Исходные данные:')
    doc.add_paragraph('• Численность рабочих: 200 человек', style='List Bullet')
    doc.add_paragraph('• Масштаб: L (Крупный)', style='List Bullet')
    doc.add_paragraph('• Автотранспорт: 15 единиц', style='List Bullet')
    doc.add_paragraph('• Иностранные рабочие: 80 человек', style='List Bullet')

    doc.add_paragraph('\nРасчёт обязательных должностей:')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    headers = ['Должность', 'Формула', 'Расчёт', 'Результат']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'D9E2F3')

    calc_mandatory = [
        ('Руководитель проекта', '1', '—', '1'),
        ('Производитель работ', 'ceil(200/25)', '8.0 → 8', '8'),
        ('Мастер', 'ceil(200/15)', '13.3 → 14', '14'),
        ('Специалист по ОТ', 'ceil(200/50)', '4.0 → 4', '4'),
        ('Специалист по общ. вопр.', 'ceil(28/15)', '1.9 → 2', '2'),
        ('Кладовщик', 'ceil(200/50)', '4.0 → 4', '4'),
    ]

    for row_data in calc_mandatory:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    doc.add_paragraph('\nРасчёт условных должностей:')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'FFEB9C')

    calc_conditional = [
        ('Водитель/Механик', 'ceil(15/3)', '5.0 → 5', '5'),
        ('Спец. по сопровождению', 'ceil(80/100)', '0.8 → 1', '1'),
    ]

    for row_data in calc_conditional:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    p = doc.add_paragraph('\nИТОГО: ')
    p.add_run('33 + 6 = 39 человек ИТР').bold = True

    doc.add_page_break()

    # === ПРИЛОЖЕНИЕ А ===
    doc.add_heading('Приложение А. Условные вопросы калькулятора', level=1)

    doc.add_paragraph(
        'При использовании калькулятора пользователю задаются следующие вопросы '
        'для определения необходимости условных должностей:'
    )

    questions = [
        ('1. Автотранспорт',
         'Используется ли автотранспорт на проекте?',
         'Если да: укажите количество единиц техники'),
        ('2. Строительные леса',
         'Используются ли строительные леса?',
         'Если да: укажите общую площадь лесов (м²)'),
        ('3. Иностранные рабочие',
         'Привлекаются ли иностранные рабочие?',
         'Если да: укажите количество человек и страну происхождения'),
        ('4. Охрана объекта',
         'Требуется ли охрана объекта?',
         'Если да: укажите количество постов и режим (круглосуточно/дневной)'),
        ('5. Проектные работы',
         'Выполняются ли проектные работы на площадке?',
         'При положительном ответе включается инженер-конструктор'),
    ]

    for title, question, note in questions:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        p.add_run('Вопрос: ').bold = True
        p.add_run(question)
        p = doc.add_paragraph()
        p.add_run('Примечание: ').italic = True
        p.add_run(note).italic = True

    # === СОХРАНЕНИЕ ===
    output_path = '/home/user/itr-calculator-reakt/docs/МЕТОДИКА_РАСЧЕТА_ИТР_v2.docx'

    # Создаём директорию если не существует
    import os
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc.save(output_path)
    print(f'Документ сохранён: {output_path}')
    return output_path

if __name__ == '__main__':
    create_methodology_document()
